"""
Exports generated questions to a formatted Word (.docx) file.

Layout (per-section grouped exam paper):
  Section 1: Multiple Choice   → numbered with A/B/C/D options (أ/ب/ج/د for Arabic)
  Section 2: True or False     → numbered statements with answer bracket
  Section 3: Short Answer      → numbered with ruled blank lines
  Section 4: Fill in the Blank → numbered sentences with blank inside
  Section 5: Essay             → numbered with ruled blank lines

Answer Key appended on a new page when include_answers=True.
Supports Arabic (RTL) and English (LTR) documents.
"""

import re
import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Fonts ─────────────────────────────────────────────────────────────────────
FONT_LATIN  = "Calibri"
FONT_ARABIC = "Arial"   # widely available, good Arabic glyphs

# ── Colours ───────────────────────────────────────────────────────────────────
COL_TITLE     = RGBColor(0x1F, 0x4E, 0x79)   # dark navy
COL_SECTION   = RGBColor(0x1F, 0x4E, 0x79)   # dark navy
COL_INST      = RGBColor(0x50, 0x50, 0x50)   # dark grey
COL_META      = RGBColor(0x70, 0x70, 0x70)   # medium grey
COL_CORRECT   = RGBColor(0x00, 0x70, 0x00)   # green (answer key)
COL_EXPLAIN   = RGBColor(0x60, 0x60, 0x60)   # grey (explanation)
COL_BLANK     = RGBColor(0xBB, 0xBB, 0xBB)   # light grey (ruled lines)


# ── Field normalisers ──────────────────────────────────────────────────────────

def _normalize_options(options) -> dict:
    """Accept both dict {'A': '...'} and list [{'text': ...}] option formats."""
    if isinstance(options, dict):
        return options
    if isinstance(options, list):
        result = {}
        for i, opt in enumerate(options[:4]):
            key = ["A", "B", "C", "D"][i]
            if isinstance(opt, dict):
                text = (opt.get("text") or opt.get("option") or
                        opt.get("value") or opt.get("content") or str(opt))
            else:
                text = str(opt)
            result[key] = text
        return result
    return {}


def _get_correct_answer(q: dict) -> str:
    return str(q.get("correct_answer") or q.get("answer") or "").strip()


def _get_model_answer(q: dict) -> str:
    return str(
        q.get("model_answer") or q.get("answer") or q.get("correct_answer") or ""
    ).strip()


def _get_question_text(q: dict) -> str:
    return (q.get("question") or q.get("statement") or "").strip()


def _strip_option_prefix(text: str) -> str:
    """Remove leading 'A.', 'B)', '(C)' etc. that the LLM sometimes adds to option text."""
    return re.sub(r'^[\(\[]?[A-Da-d][\.\)\]:\-]\s*', '', text.strip()).strip()


# ── RTL helpers ───────────────────────────────────────────────────────────────

def _set_para_rtl(paragraph):
    """Force a paragraph to right-to-left (Arabic) direction."""
    pPr = paragraph._p.get_or_add_pPr()
    # bidi must be first child
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    # right-align
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    existing_jc = pPr.find(qn('w:jc'))
    if existing_jc is not None:
        pPr.remove(existing_jc)
    pPr.append(jc)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_doc_rtl(doc):
    """Mark the whole document as RTL (affects default text direction)."""
    settings = doc.settings.element
    bidi_el = OxmlElement('w:bidi')
    bidi_el.set(qn('w:val'), '1')
    settings.append(bidi_el)


# ── Run font helper ───────────────────────────────────────────────────────────

def _fmt(run, size_pt=None, bold=False, italic=False,
         color: RGBColor = None, is_arabic: bool = False):
    """Apply font, size, weight, and colour to a run."""
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    # Set explicit fonts (Latin + Complex Script for Arabic)
    rPr = run._r.get_or_add_rPr()
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_LATIN)
    rFonts.set(qn('w:hAnsi'), FONT_LATIN)
    rFonts.set(qn('w:cs'), FONT_ARABIC if is_arabic else FONT_LATIN)
    rPr.insert(0, rFonts)


# ── Border helpers ────────────────────────────────────────────────────────────

def _add_bottom_border(paragraph, color: str = "CCCCCC", size: str = "4"):
    """Add a thin bottom border to a paragraph (used as visual separator)."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing pBdr if any
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _blank_line(doc, is_arabic: bool):
    """A single ruled answer line (paragraph with a light bottom border)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    _add_bottom_border(p, color="CCCCCC", size="4")
    if is_arabic:
        _set_para_rtl(p)
    return p


# ── Labels & lookup tables ────────────────────────────────────────────────────

_TYPE_LABELS_AR = {
    "mcq":          "اختيار من متعدد",
    "true_false":   "صح أم خطأ",
    "short_answer": "إجابة قصيرة",
    "fill_blank":   "ملء الفراغ",
    "essay":        "سؤال مقالي",
}
_TYPE_LABELS_EN = {
    "mcq":          "Multiple Choice",
    "true_false":   "True or False",
    "short_answer": "Short Answer",
    "fill_blank":   "Fill in the Blank",
    "essay":        "Essay",
}

_INSTRUCTIONS_AR = {
    "mcq":          "اختر الإجابة الصحيحة من بين البدائل التالية.",
    "true_false":   "ضع علامة (صح) أو (خطأ) في المربع أمام كل عبارة.",
    "short_answer": "أجب عن الأسئلة التالية بإجابات قصيرة.",
    "fill_blank":   "أكمل الجمل الآتية بالكلمة أو العبارة المناسبة.",
    "essay":        "أجب عن الأسئلة التالية إجابة وافية.",
}
_INSTRUCTIONS_EN = {
    "mcq":          "Choose the correct answer from the options below.",
    "true_false":   "Write True or False in the box next to each statement.",
    "short_answer": "Answer the following questions briefly.",
    "fill_blank":   "Fill in the blanks with the appropriate word or phrase.",
    "essay":        "Answer the following questions in detail.",
}

_ORDINALS_AR  = ["الأول", "الثاني", "الثالث", "الرابع", "الخامس"]
_TYPE_ORDER   = ["mcq", "true_false", "short_answer", "fill_blank", "essay"]
_OPT_EN       = ["A", "B", "C", "D"]
_OPT_AR       = ["أ", "ب", "ج", "د"]


def _section_header_text(q_type: str, section_num: int, count: int, is_arabic: bool) -> str:
    if is_arabic:
        label   = _TYPE_LABELS_AR.get(q_type, q_type)
        ordinal = _ORDINALS_AR[section_num - 1] if section_num <= 5 else f"رقم {section_num}"
        return f"السؤال {ordinal}: {label}  ({count})"
    label = _TYPE_LABELS_EN.get(q_type, q_type)
    return f"Section {section_num}: {label}  ({count} questions)"


# ── Question renderer ─────────────────────────────────────────────────────────

def _render_question_item(doc, q: dict, num: int, q_type: str, is_arabic: bool):
    """Render a single numbered question with its sub-elements."""
    text = _get_question_text(q)
    if not text:
        return   # skip empty/failed questions

    # ── Numbered question line ────────────────────────────────────────────────
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_before = Pt(10)
    q_para.paragraph_format.space_after  = Pt(3)

    if q_type == "true_false":
        # Add a small answer box [ ] for the student to tick
        box = "[ ]"
        line = f"{num}.  {box}  {text}" if not is_arabic else f"{text}  {box}  .{num}"
    else:
        line = f"{num}.  {text}"

    q_run = q_para.add_run(line)
    _fmt(q_run, size_pt=11, is_arabic=is_arabic)
    if is_arabic:
        _set_para_rtl(q_para)

    # ── Type-specific sub-content ─────────────────────────────────────────────
    if q_type == "mcq":
        options = _normalize_options(q.get("options", {}))
        letters = _OPT_AR if is_arabic else _OPT_EN
        for j, key in enumerate(_OPT_EN):
            raw  = options.get(key, "").strip()
            text = _strip_option_prefix(raw)
            if not text:
                continue
            opt = doc.add_paragraph()
            opt.paragraph_format.space_before = Pt(2)
            opt.paragraph_format.space_after  = Pt(2)
            # Proper indent — no leading spaces
            if is_arabic:
                opt.paragraph_format.right_indent = Inches(0.35)
            else:
                opt.paragraph_format.left_indent = Inches(0.35)
            opt_run = opt.add_run(f"{letters[j]}.  {text}")
            _fmt(opt_run, size_pt=10, is_arabic=is_arabic)
            if is_arabic:
                _set_para_rtl(opt)

    elif q_type in ("short_answer", "essay"):
        lines = 3 if q_type == "short_answer" else 6
        for _ in range(lines):
            _blank_line(doc, is_arabic)

    # fill_blank and true_false: nothing extra — question carries the blank/box


# ── Answer-key renderer ───────────────────────────────────────────────────────

def _render_answer_key(doc, grouped: dict, is_arabic: bool):
    """Append an Answer Key section on a new page."""
    doc.add_page_break()

    # "Answer Key" / "مفتاح الإجابات"
    hdr_text = "مفتاح الإجابات" if is_arabic else "Answer Key"
    h = doc.add_heading(hdr_text, 1)
    h.paragraph_format.space_after = Pt(6)
    if h.runs:
        _fmt(h.runs[0], size_pt=16, bold=True, color=COL_TITLE, is_arabic=is_arabic)
    if is_arabic:
        _set_para_rtl(h)

    section_num = 0
    for q_type in _TYPE_ORDER:
        if q_type not in grouped:
            continue
        valid_qs = [q for q in grouped[q_type] if _get_question_text(q)]
        if not valid_qs:
            continue

        section_num += 1

        # Section sub-header
        sh = doc.add_paragraph()
        sh.paragraph_format.space_before = Pt(14)
        sh.paragraph_format.space_after  = Pt(4)
        sh_run = sh.add_run(_section_header_text(q_type, section_num, len(valid_qs), is_arabic))
        _fmt(sh_run, size_pt=11, bold=True, color=COL_SECTION, is_arabic=is_arabic)
        _add_bottom_border(sh, color="AAAAAA", size="4")
        if is_arabic:
            _set_para_rtl(sh)

        for i, q in enumerate(valid_qs, 1):
            ans = doc.add_paragraph()
            ans.paragraph_format.space_before = Pt(4)
            ans.paragraph_format.space_after  = Pt(4)

            if q_type == "mcq":
                correct = _get_correct_answer(q).upper()
                explanation = q.get("explanation") or q.get("reason") or ""
                if is_arabic and correct in _OPT_EN:
                    display = _OPT_AR[_OPT_EN.index(correct)]
                else:
                    display = correct
                r1 = ans.add_run(f"{i}.  ")
                _fmt(r1, size_pt=10, is_arabic=is_arabic)
                r2 = ans.add_run(display)
                _fmt(r2, size_pt=10, bold=True, color=COL_CORRECT, is_arabic=is_arabic)
                if explanation:
                    r3 = ans.add_run(f"  —  {explanation}")
                    _fmt(r3, size_pt=9, color=COL_EXPLAIN, is_arabic=is_arabic)

            elif q_type == "true_false":
                correct = _get_correct_answer(q)
                explanation = q.get("explanation") or q.get("reason") or ""
                if is_arabic:
                    display = "صح" if correct.lower() in ("true", "صح", "t", "1", "yes") else "خطأ"
                else:
                    display = "True" if correct.lower() in ("true", "t", "1", "yes") else "False"
                r1 = ans.add_run(f"{i}.  ")
                _fmt(r1, size_pt=10, is_arabic=is_arabic)
                r2 = ans.add_run(display)
                _fmt(r2, size_pt=10, bold=True, color=COL_CORRECT, is_arabic=is_arabic)
                if explanation:
                    r3 = ans.add_run(f"  —  {explanation}")
                    _fmt(r3, size_pt=9, color=COL_EXPLAIN, is_arabic=is_arabic)

            elif q_type == "short_answer":
                model = _get_model_answer(q)
                r1 = ans.add_run(f"{i}.  ")
                _fmt(r1, size_pt=10, bold=True, is_arabic=is_arabic)
                r2 = ans.add_run(model)
                _fmt(r2, size_pt=10, is_arabic=is_arabic)

            elif q_type == "fill_blank":
                correct = _get_correct_answer(q)
                r1 = ans.add_run(f"{i}.  ")
                _fmt(r1, size_pt=10, is_arabic=is_arabic)
                r2 = ans.add_run(correct)
                _fmt(r2, size_pt=10, bold=True, color=COL_CORRECT, is_arabic=is_arabic)

            elif q_type == "essay":
                key_points = q.get("key_points") or []
                r1 = ans.add_run(f"{i}.  ")
                _fmt(r1, size_pt=10, bold=True, is_arabic=is_arabic)
                bullets = "  •  ".join(str(p) for p in key_points if p)
                r2 = ans.add_run(bullets or _get_model_answer(q))
                _fmt(r2, size_pt=10, is_arabic=is_arabic)

            if is_arabic:
                _set_para_rtl(ans)


# ── Main exporter ──────────────────────────────────────────────────────────────

def export_to_docx(
    questions:       list[dict],
    title:           str  = "Exam Questions",
    document_id:     str  = "",
    difficulty:      str  = "",          # kept for API compat, not displayed
    language:        str  = "en",
    include_answers: bool = True,
) -> bytes:
    """
    Generate a polished exam-paper Word (.docx) from a list of questions.

    Returns raw .docx bytes.
    """
    is_arabic = (language == "ar")
    doc = Document()

    if is_arabic:
        _set_doc_rtl(doc)

    # ── Page margins ──────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.15)
        sec.right_margin  = Inches(1.15)

    # ── Title ─────────────────────────────────────────────────────────────────
    t = doc.add_heading(title, 0)
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after  = Pt(6)
    if t.runs:
        _fmt(t.runs[0], size_pt=20, bold=True, color=COL_TITLE, is_arabic=is_arabic)
    if is_arabic:
        _set_para_rtl(t)

    # ── Student info row ──────────────────────────────────────────────────────
    today = datetime.date.today().strftime("%Y-%m-%d")
    if is_arabic:
        info_line = f"الاسم: ________________________    الصف: ___________    التاريخ: {today}    الدرجة: _______ / "
    else:
        info_line = f"Name: ________________________    Class: ___________    Date: {today}    Grade: _______ / "

    info = doc.add_paragraph()
    info.paragraph_format.space_before = Pt(4)
    info.paragraph_format.space_after  = Pt(4)
    info_run = info.add_run(info_line)
    _fmt(info_run, size_pt=10, is_arabic=is_arabic)
    _add_bottom_border(info, color="AAAAAA", size="6")
    if is_arabic:
        _set_para_rtl(info)

    # ── Total-question count ──────────────────────────────────────────────────
    valid_total = sum(1 for q in questions if _get_question_text(q))
    if is_arabic:
        count_text = f"عدد الأسئلة الكلي: {valid_total}"
    else:
        count_text = f"Total questions: {valid_total}"

    count_p = doc.add_paragraph()
    count_p.paragraph_format.space_before = Pt(6)
    count_p.paragraph_format.space_after  = Pt(14)
    count_run = count_p.add_run(count_text)
    _fmt(count_run, size_pt=9, italic=True, color=COL_META, is_arabic=is_arabic)
    if is_arabic:
        _set_para_rtl(count_p)

    # ── Group questions by type ───────────────────────────────────────────────
    grouped: dict[str, list[dict]] = {}
    for q in questions:
        grouped.setdefault(q.get("type", "other"), []).append(q)

    # ── Render sections ───────────────────────────────────────────────────────
    section_num = 0
    for q_type in _TYPE_ORDER:
        if q_type not in grouped:
            continue
        valid_qs = [q for q in grouped[q_type] if _get_question_text(q)]
        if not valid_qs:
            continue

        section_num += 1

        # Section header with blue bottom border
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(18)
        hdr.paragraph_format.space_after  = Pt(2)
        hdr_run = hdr.add_run(_section_header_text(q_type, section_num, len(valid_qs), is_arabic))
        _fmt(hdr_run, size_pt=13, bold=True, color=COL_SECTION, is_arabic=is_arabic)
        _add_bottom_border(hdr, color="2E74B5", size="8")
        if is_arabic:
            _set_para_rtl(hdr)

        # Instructions line (italic, grey)
        instr = (_INSTRUCTIONS_AR if is_arabic else _INSTRUCTIONS_EN).get(q_type, "")
        if instr:
            ip = doc.add_paragraph()
            ip.paragraph_format.space_before = Pt(4)
            ip.paragraph_format.space_after  = Pt(8)
            ir = ip.add_run(instr)
            _fmt(ir, size_pt=10, italic=True, color=COL_INST, is_arabic=is_arabic)
            if is_arabic:
                _set_para_rtl(ip)

        # Questions
        for i, q in enumerate(valid_qs, 1):
            _render_question_item(doc, q, i, q_type, is_arabic)

    # ── Answer key ────────────────────────────────────────────────────────────
    if include_answers and grouped:
        _render_answer_key(doc, grouped, is_arabic)

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
